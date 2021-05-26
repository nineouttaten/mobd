import datetime

import docx
import pandas as pd
import sqlite3
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5 import uic, QtWidgets
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, QMessageBox
from dateutil.relativedelta import relativedelta

Form, Window = uic.loadUiType("give.ui")
from main import main_Ui
from math import ceil


# def make_document(matrix, first_name, last_name, patronymic, series, number, months):
#     person = first_name + " " + last_name + " " + patronymic
#     person_key = series + " " + number
#     doc = docx.Document()
#     doc.add_heading("График платежей по кредиту\nКлиент:", person).alignment = 1
#     head = ['Дата платежа', 'Сумма долга', 'Проценты', 'Основная часть', 'Сумма долга на окончание периода']
#     table = doc.add_table(rows=months + 1, cols=6)
#     for col in range(6):
#         cell = table.cell(0, col)
#         cell.text = head[col]
#     for row in range(months):
#         for col in range(6):
#             cell = table.cell(row, col)
#             cell.text = matrix[row, col]
#     table.style = "Table Grid"
#     doc.save(person_key + ".docx")


class Give_Ui(QtWidgets.QDialog, Form):
    global clients
    clients = sqlite3.connect("clients.db")
    global pay
    pay = sqlite3.connect("pay.db")
    global clients_cursor
    clients_cursor = clients.cursor()
    global pay_cursor
    pay_cursor = pay.cursor()

    def __init__(self):
        super(Give_Ui, self).__init__()
        self.setupUi(self)
        self.btnHome.clicked.connect(self.btnHomePressed)
        self.btnGenerate.clicked.connect(self.btnGeneratePressed)
        self.btnGive.clicked.connect(self.btnGivePressed)
        clients_cursor.execute("""CREATE TABLE IF NOT EXISTS clients (
            last_name TEXT,
            first_name TEXT,
            patronymic TEXT,
            birthday DATE,
            series TEXT,
            number TEXT,
            key TEXT,
            passport_date DATE,
            passport_office TEXT,
            city TEXT,
            street TEXT,
            house TEXT,
            credit INTEGER,
            procent INTEGER,
            insurance BOOL,
            months INTEGER
        )""")
        pay_cursor.execute("""CREATE TABLE IF NOT EXISTS pay (
                    key TEXT,
                    payment_date TEXT,
                    sum INTEGER
                )""")
        clients.commit()

    def btnHomePressed(self):
        self.window = main_Ui()
        self.window.show()
        self.close()

    def btnGeneratePressed(self):

        information = pd.Series(self.get_information(),
                                index=["last_name", "first_name", "patronymic", "birthday", "series", "number",
                                       "passport_date", "passport_office", "city", "street", "house", "credit",
                                       "procent", "insurance", "months"])
        print(information)

    # Вычисление платежа по сумме кредита
    def graphics(self, credit, procent, insurance, months=24):
        # Вычисление доли процентов в ежемесячном взносе
        print(insurance)
        print(int(insurance))
        I = (procent + int(insurance)) / 12 / 100
        # Вычисление размера ежемесячного платежа
        x = ceil(credit * (I + I / ((1 + I) ** months - 1)))
        return [I, x]

    # Создание матрицы платежей
    def payment_matrix(self, credit, payment, I, months):
        matrix = []
        S = credit
        new_date = datetime.datetime.today().date()
        for i in range(months - 1):
            new_date = new_date + relativedelta(months=+1)
            pc = int(S * I)
            main = payment - pc
            result_S = S - main
            matrix.append([new_date, S, payment, pc, main, result_S])
            S -= (payment - pc)
        new_date = new_date + relativedelta(months=+1)
        matrix.append([new_date, S, S, 0, S, 0])
        matrix = pd.DataFrame(matrix, index=[list(range(1, months + 1))],
                              columns=["Date", "Sum", 'payment', "percent", "body", "Sum_after"])
        return matrix

    def get_information(self):
        try:
            last_name = self.lineEdit.text()
            first_name = self.lineEdit_2.text()
            patronymic = self.lineEdit_3.text()
            birthday = self.dateEdit_2.text()
            series = self.lineEdit_4.text()
            number = self.lineEdit_5.text()
            key = series + number
            passport_date = self.dateEdit.text()
            passport_office = self.textEdit.toPlainText()
            city = self.lineEdit_8.text()
            street = self.lineEdit_9.text()
            house = self.lineEdit_10.text()
            credit = int(self.lineEdit_6.text())
            procent = int(self.lineEdit_7.text())
            insurance = self.checkBox.isChecked()
            months = int(self.lineEdit_11.text())
            inf = [last_name, first_name, patronymic, birthday, series, number, key, passport_date, passport_office, city,
                   street, house, credit, procent, insurance, months]
            correct = self.isCorrect(inf)
            if not correct:
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Critical)
                msg.setText("Введены не все данные")
                msg.setWindowTitle("Ошибка")
                msg.exec()
                return False
        #if correct:
            return inf
        #else:
            # ToDo: Окно об ошибке ввода данных
        except Exception:
            print("xd")
            return False

    def isCorrect(self, information):
        check = True
        for value in information:
            if value == "":
                check = False
        str_birthday = information[3]
        birthday = datetime.datetime.strptime(str_birthday, "%d.%m.%Y").date()
        after_18 = birthday + relativedelta(years=+18)
        if datetime.date.today() < after_18:
            check = False
        return check

    def btnGivePressed(self):
        if not self.get_information():
            if not self.lineEdit_6.text().isdigit() or not self.lineEdit_7.text().isdigit() or not self.lineEdit_11.text().isdigit():
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Critical)
                msg.setText("Некоректно введены числовые значения")
                msg.setWindowTitle("Ошибка")
                msg.exec()

        else:
            information = pd.Series(self.get_information(),
                                    index=["last_name", "first_name", "patronymic", "birthday", "series", "number", "key",
                                           "passport_date", "passport_office", "city", "street", "house", "credit",
                                           "procent", "insurance", "months"])
            clients_cursor.execute(f"INSERT INTO clients VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", (
            information["last_name"], information["first_name"], information["patronymic"], information["birthday"],
            information["series"], information["number"], information["key"], information["passport_date"],
            information["passport_office"], information["city"], information["street"], information["house"],
            information["credit"], information["procent"], information["insurance"], information["months"],))
            clients.commit()
            dop_info = self.graphics(information["credit"], information["procent"], information["insurance"],
                                     information["months"])
            matrix = self.payment_matrix(information["credit"], dop_info[1], dop_info[0], information["months"])
            for i in range(information["months"]):
                pay_cursor.execute(f"INSERT INTO pay VALUES(?, ?, ?)", (
                information["series"] + information["number"], matrix.iloc[i, :]["Date"],
                int(matrix.iloc[i, :]["payment"])))
                pay.commit()

            for i in pay_cursor.execute("SELECT * FROM pay"):
                print(i)


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    w = Give_Ui()
    w.setWindowTitle('Выдача кредита МФО «‎Купи не Копи»‎')
    w.show()  # show window
    sys.exit(app.exec_())
